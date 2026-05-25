from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

path = Path("temp_section_break.docx")
doc = Document()
doc.add_heading("1. Aim", level=1)
doc.add_paragraph("Paragraph 1 of section 1.")
doc.add_section(WD_SECTION.NEW_PAGE)
doc.add_heading("2. Theory", level=1)
doc.add_paragraph("Paragraph 1 of section 2.")
doc.save(path)
print("saved", path)
for idx, para in enumerate(Document(path).paragraphs):
    pPr = para._p.find(qn("w:pPr"))
    sb = pPr.find(qn("w:sectPr")) if pPr is not None else None
    print(idx, repr(para.text), "runs=", len(para.runs), "sectPr=", sb is not None)
    if idx == 1:
        print("xml:", para._p.xml)
