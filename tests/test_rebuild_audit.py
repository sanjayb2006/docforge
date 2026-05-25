"""
tests/test_rebuild_audit.py

Step 0 — Rebuild Audit Tests.

Tests the rebuild_audit module directly (no DB, no FastAPI, no async).
Creates real DOCX fixtures, runs rebuild, runs audit, asserts findings.

Also includes a standalone runner that prints a human-readable audit
report — useful for eyeballing what the current rebuilder actually does.

Run all tests:
    pytest tests/test_rebuild_audit.py -v

Run the human-readable report for one specific scenario:
    python tests/test_rebuild_audit.py report

Run the full matrix report (all scenarios):
    python tests/test_rebuild_audit.py matrix
"""

from __future__ import annotations

import io
import json
import sys
import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from app.services.docx.parser import parse_docx
from app.services.docx.style_extractor import extract_style_profile
from app.services.docx.rebuilder import rebuild_docx
from app.services.docx.rebuild_audit import (
    snapshot_document,
    diff_snapshots,
    log_audit,
    audit_to_dict,
    RebuildAudit,
)


# ══════════════════════════════════════════════════════════════════════════════
# FIXTURE BUILDERS
# ══════════════════════════════════════════════════════════════════════════════

def _minimal_png() -> bytes:
    """1×1 red PNG for image tests."""
    import base64
    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8"
        "z8BQDwADhQGAWjR9awAAAABJRU5ErkJggg=="
    )


def build_lab_report(path: Path) -> None:
    """Standard VTU lab report — baseline fixture."""
    doc = DocxDocument()
    doc.add_heading("BCSL657D — DevOps Lab", level=0)
    for sec, body in [
        ("1. Aim",        "To implement a CI/CD pipeline using Jenkins and Docker."),
        ("2. Theory",     "Continuous Integration is a software engineering practice."),
        ("3. Procedure",  "Install Jenkins. Configure pipeline. Run build."),
        ("4. Result",     "Pipeline executed successfully with all stages passing."),
        ("5. Conclusion", "The experiment demonstrated the value of CI/CD automation."),
    ]:
        doc.add_heading(sec, level=1)
        doc.add_paragraph(body)
    doc.save(str(path))


def build_mixed_format(path: Path) -> None:
    """Document with bold/italic runs, custom spacing, nested headings."""
    doc = DocxDocument()
    doc.add_heading("Mixed Format Report", level=0)

    doc.add_heading("1. Introduction", level=1)
    para = doc.add_paragraph()
    para.add_run("Normal text. ")
    bold_run = para.add_run("Bold section.")
    bold_run.bold = True
    italic_run = para.add_run(" Italic note.")
    italic_run.italic = True

    doc.add_heading("1.1 Background", level=2)
    p = doc.add_paragraph("Indented paragraph with custom spacing.")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Pt(36)

    doc.add_heading("2. Data", level=1)
    doc.add_paragraph("Standard body paragraph here.")
    doc.save(str(path))


def build_with_table(path: Path) -> None:
    """Document containing a table."""
    doc = DocxDocument()
    doc.add_heading("Report with Table", level=1)
    doc.add_paragraph("See data below.")
    t = doc.add_table(rows=3, cols=3)
    for i, row in enumerate(t.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"R{i}C{j}"
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("Table data reviewed.")
    doc.save(str(path))


def build_with_image(path: Path) -> None:
    """Document containing an inline image."""
    doc = DocxDocument()
    doc.add_heading("Report with Image", level=1)
    doc.add_paragraph("Figure below:")
    doc.add_picture(io.BytesIO(_minimal_png()), width=Inches(1))
    doc.add_paragraph("Caption: Test image.")
    doc.add_heading("Result", level=1)
    doc.add_paragraph("Image rendered correctly.")
    doc.save(str(path))


def build_large(path: Path, n_sections: int = 20) -> None:
    """Large document to test count stability."""
    doc = DocxDocument()
    doc.add_heading("Large Report", level=0)
    for i in range(n_sections):
        doc.add_heading(f"Section {i+1}", level=1)
        doc.add_paragraph(f"Content for section {i+1}. " * 5)
        if i % 3 == 0:
            doc.add_heading(f"Section {i+1}.1", level=2)
            doc.add_paragraph("Sub-section body text.")
    doc.save(str(path))


# ══════════════════════════════════════════════════════════════════════════════
# TEST HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _run_pipeline(build_fn, ai_results: dict | None = None) -> tuple[Path, Path, RebuildAudit]:
    """
    Build a DOCX with build_fn, run rebuild, run audit.
    Returns (original_path, output_path, audit).
    Caller is responsible for cleanup.
    """
    orig_tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    orig_path = Path(orig_tmp.name)
    orig_tmp.close()
    build_fn(orig_path)

    out_tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    out_path = Path(out_tmp.name)
    out_tmp.close()
    out_path.unlink()

    structure     = parse_docx(orig_path)
    style_profile = extract_style_profile(orig_path)

    rebuild_docx(
        original_path=orig_path,
        output_path=out_path,
        structure=structure,
        ai_results=ai_results or {},
        style_profile=style_profile,
    )

    orig_snap = snapshot_document(orig_path, label="original")
    out_snap  = snapshot_document(out_path,  label="output")
    audit     = diff_snapshots(orig_snap, out_snap, ai_sections=set((ai_results or {}).keys()))

    return orig_path, out_path, audit


def _cleanup(*paths: Path) -> None:
    for p in paths:
        p.unlink(missing_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
# TESTS — AUDIT MODULE CORRECTNESS
# ══════════════════════════════════════════════════════════════════════════════

class TestSnapshotModule:
    """snapshot_document() and diff_snapshots() behave correctly."""

    def test_snapshot_returns_correct_counts(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = Path(f.name)
        build_lab_report(path)
        snap = snapshot_document(path, label="test")
        assert snap.para_count > 0
        assert snap.heading_count >= 5     # title + 5 h1s
        assert len(snap.heading_texts) >= 5
        assert snap.table_count == 0
        path.unlink(missing_ok=True)

    def test_identical_docs_zero_delta(self):
        """Snapshotting the same file twice → zero paragraph delta."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = Path(f.name)
        build_lab_report(path)
        s1 = snapshot_document(path, "a")
        s2 = snapshot_document(path, "b")
        audit = diff_snapshots(s1, s2)
        assert audit.para_count_delta == 0
        assert audit.tables_lost == 0
        assert audit.headings_missing == []
        path.unlink(missing_ok=True)

    def test_audit_to_dict_is_serialisable(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = Path(f.name)
        build_lab_report(path)
        s = snapshot_document(path)
        audit = diff_snapshots(s, s)
        d = audit_to_dict(audit)
        json_str = json.dumps(d)  # must not raise
        assert "severity" in json_str
        path.unlink(missing_ok=True)

    def test_missing_heading_detected(self):
        """If output is missing a heading, headings_missing is non-empty."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            orig = Path(f.name)
        build_lab_report(orig)
        orig_snap = snapshot_document(orig)

        # Build a reduced output with fewer headings
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            out = Path(f.name)
        d = DocxDocument()
        d.add_heading("1. Aim", level=1)
        d.add_paragraph("Only one section.")
        d.save(str(out))
        out_snap = snapshot_document(out)

        audit = diff_snapshots(orig_snap, out_snap)
        assert len(audit.headings_missing) > 0
        assert audit.severity in ("major", "critical")

        orig.unlink(missing_ok=True)
        out.unlink(missing_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
# TESTS — REBUILD AUDIT SCENARIOS
# ══════════════════════════════════════════════════════════════════════════════

class TestBasicRebuildAudit:
    """Audit baseline lab report rebuild (no AI replacements)."""

    def setup_method(self):
        orig, out, self.audit = _run_pipeline(build_lab_report, ai_results={})
        self.orig, self.out = orig, out

    def teardown_method(self):
        _cleanup(self.orig, self.out)

    def test_all_headings_preserved(self):
        assert self.audit.headings_missing == [], (
            f"Lost headings: {self.audit.headings_missing}"
        )

    def test_no_tables_lost(self):
        assert self.audit.tables_lost == 0

    def test_no_images_lost(self):
        assert self.audit.images_lost == 0

    def test_no_text_duplication(self):
        assert self.audit.text_duplications == 0, (
            f"Duplicated text detected in {self.audit.text_duplications} paragraphs"
        )

    def test_paragraph_count_reasonable(self):
        # Allow ±3 paragraphs for empty paragraph handling differences
        assert abs(self.audit.para_count_delta) <= 3, (
            f"Paragraph count delta too large: {self.audit.para_count_delta}"
        )

    def test_severity_not_critical(self):
        assert self.audit.severity != "critical", (
            f"Severity is critical. Full audit:\n{json.dumps(audit_to_dict(self.audit), indent=2)}"
        )


class TestAiSectionRebuildAudit:
    """Audit rebuild with AI content injected into some sections."""

    def setup_method(self):
        ai = {
            "1. Aim":    "AI-generated aim: implement Maven build pipeline.",
            "4. Result": "AI-generated result: build completed in 3.2 seconds.",
        }
        orig, out, self.audit = _run_pipeline(build_lab_report, ai_results=ai)
        self.orig, self.out, self.ai = orig, out, ai

    def teardown_method(self):
        _cleanup(self.orig, self.out)

    def test_all_headings_still_present(self):
        assert self.audit.headings_missing == [], (
            f"Headings lost after AI injection: {self.audit.headings_missing}"
        )

    def test_no_text_duplication(self):
        assert self.audit.text_duplications == 0, (
            f"Duplication found after AI injection"
        )

    def test_no_text_truncation_in_preserved_sections(self):
        # Theory, Procedure, Conclusion were NOT AI-replaced — they must be intact
        assert self.audit.text_truncations == 0, (
            f"Text truncated in {self.audit.text_truncations} paragraphs"
        )


class TestMixedFormatAudit:
    """Audit rebuild of document with bold/italic runs and custom spacing."""

    def setup_method(self):
        orig, out, self.audit = _run_pipeline(build_mixed_format, ai_results={})
        self.orig, self.out = orig, out

    def teardown_method(self):
        _cleanup(self.orig, self.out)

    def test_headings_preserved(self):
        assert self.audit.headings_missing == []

    def test_no_duplication(self):
        assert self.audit.text_duplications == 0


class TestTableRebuildAudit:
    """Tables must survive the rebuild cycle."""

    def setup_method(self):
        orig, out, self.audit = _run_pipeline(build_with_table, ai_results={})
        self.orig, self.out = orig, out

    def teardown_method(self):
        _cleanup(self.orig, self.out)

    def test_table_not_lost(self):
        assert self.audit.tables_lost == 0, (
            f"Table lost in rebuild. orig={self.audit.orig_table_count} "
            f"out={self.audit.out_table_count}"
        )

    def test_table_count_matches(self):
        assert self.audit.orig_table_count == self.audit.out_table_count


class TestImageRebuildAudit:
    """Images must survive the rebuild cycle."""

    def setup_method(self):
        orig, out, self.audit = _run_pipeline(build_with_image, ai_results={})
        self.orig, self.out = orig, out

    def teardown_method(self):
        _cleanup(self.orig, self.out)

    def test_image_not_lost(self):
        assert self.audit.images_lost == 0, (
            f"Image lost in rebuild. orig={self.audit.orig_image_count} "
            f"out={self.audit.out_image_count}"
        )

    def test_image_count_preserved(self):
        assert self.audit.orig_image_count == self.audit.out_image_count

    def test_headings_preserved(self):
        # Image paragraph shifts index alignment — but headings must survive
        assert self.audit.headings_missing == [], f"Headings lost: {self.audit.headings_missing}"

    def test_no_tables_lost(self):
        assert self.audit.tables_lost == 0


class TestLargeDocRebuildAudit:
    """Large documents should maintain structural stability."""

    def setup_method(self):
        orig, out, self.audit = _run_pipeline(
            lambda p: build_large(p, n_sections=15),
            ai_results={},
        )
        self.orig, self.out = orig, out

    def teardown_method(self):
        _cleanup(self.orig, self.out)

    def test_no_headings_lost(self):
        assert self.audit.headings_missing == [], (
            f"Lost: {self.audit.headings_missing}"
        )

    def test_para_delta_stable(self):
        assert abs(self.audit.para_count_delta) <= 5, (
            f"Large doc paragraph delta: {self.audit.para_count_delta}"
        )

    def test_no_duplication(self):
        assert self.audit.text_duplications == 0


# ══════════════════════════════════════════════════════════════════════════════
# STANDALONE HUMAN-READABLE REPORTER
# ══════════════════════════════════════════════════════════════════════════════

SCENARIOS = {
    "basic":        (build_lab_report,   {}),
    "ai_injected":  (build_lab_report,   {"1. Aim": "AI aim content.", "4. Result": "AI result content."}),
    "mixed_format": (build_mixed_format, {}),
    "table":        (build_with_table,   {}),
    "image":        (build_with_image,   {}),
    "large":        (lambda p: build_large(p, 10), {}),
}


def _print_report(name: str, audit: RebuildAudit) -> None:
    sev_colors = {
        "ok":       "\033[92m",  # green
        "minor":    "\033[93m",  # yellow
        "major":    "\033[91m",  # red
        "critical": "\033[1;91m",  # bold red
    }
    reset = "\033[0m"
    sev_color = sev_colors.get(audit.severity, "")

    print(f"\n{'═'*60}")
    print(f"  SCENARIO: {name.upper()}")
    print(f"  SEVERITY: {sev_color}{audit.severity.upper()}{reset}")
    print(f"{'═'*60}")
    print(f"  Paragraphs:  {audit.orig_para_count} → {audit.out_para_count}  "
          f"(delta={audit.para_count_delta:+d})")
    print(f"  Tables:      {audit.orig_table_count} → {audit.out_table_count}  "
          f"(lost={audit.tables_lost})")
    print(f"  Images:      {audit.orig_image_count} → {audit.out_image_count}  "
          f"(lost={audit.images_lost})")
    print(f"  Headings:    {audit.orig_heading_count} → {audit.out_heading_count}  "
          f"(matched={audit.headings_matched})")
    print(f"  Runs:        {audit.orig_run_count} → {audit.out_run_count}")

    if audit.headings_missing:
        print(f"\n  ⚠  HEADINGS MISSING:")
        for h in audit.headings_missing:
            print(f"      - '{h}'")
    if audit.headings_added:
        print(f"\n  ⚠  HEADINGS ADDED (unexpected):")
        for h in audit.headings_added:
            print(f"      + '{h}'")
    if audit.margin_drifts:
        print(f"\n  ⚠  MARGIN DRIFT:")
        for side, (o, r) in audit.margin_drifts.items():
            print(f"      {side}: {o}pt → {r}pt  (Δ{r-o:+.1f}pt)")

    print(f"\n  Anomalies:")
    print(f"    style_mismatches  : {audit.style_mismatches}")
    print(f"    spacing_anomalies : {audit.spacing_anomalies}")
    print(f"    run_count_changes : {audit.run_count_anomalies}")
    print(f"    text_truncations  : {audit.text_truncations}")
    print(f"    text_duplications : {audit.text_duplications}")
    print(f"    font_changes      : {audit.font_changes}")
    print(f"    field_codes_lost  : {audit.field_codes_lost}")

    if audit.para_diffs:
        print(f"\n  Per-paragraph changes ({len(audit.para_diffs)} paragraphs affected):")
        for diff in audit.para_diffs[:8]:
            parts = []
            if diff.image_lost:        parts.append("IMAGE-LOST")
            if diff.text_duplicated:   parts.append(f"DUPLICATION(orig={diff.orig_text_len} out={diff.out_text_len})")
            if diff.text_truncated:    parts.append(f"TRUNCATION(orig={diff.orig_text_len} out={diff.out_text_len})")
            if diff.field_code_lost:   parts.append("FIELD-CODE-LOST")
            if diff.style_changed:     parts.append(f"STYLE({diff.orig_style!r}→{diff.out_style!r})")
            if diff.spacing_changed:   parts.append("SPACING")
            if diff.alignment_changed: parts.append("ALIGNMENT")
            if diff.run_count_changed: parts.append(f"RUNS({diff.orig_runs}→{diff.out_runs})")
            if diff.font_changed:      parts.append(f"FONT({diff.orig_fonts}→{diff.out_fonts})")
            print(f"    para[{diff.index:3d}]: " + "  ".join(parts))


def run_report(scenario_name: str = "basic") -> None:
    """Run one scenario and print human-readable audit report."""
    if scenario_name not in SCENARIOS:
        print(f"Unknown scenario '{scenario_name}'. Available: {list(SCENARIOS)}")
        return

    build_fn, ai = SCENARIOS[scenario_name]
    orig, out, audit = _run_pipeline(build_fn, ai_results=ai)
    _print_report(scenario_name, audit)
    _cleanup(orig, out)


def run_matrix() -> None:
    """Run all scenarios and print A/B comparison matrix: rebuilder vs transformer."""
    SEV = {"ok": "✓ ok", "minor": "~ min", "major": "! MAJ", "critical": "✗ CRIT"}

    col = f"  {'Scenario':<16} {'Engine':<13} {'Sev':<9} {'ΔPara':>6} {'ΔImg':>5} {'MissH':>6} {'Dup':>4} {'Trunc':>6} {'Style':>6}"
    width = len(col)

    print("\n" + "═" * width)
    print("  DOCFORGE — REBUILD AUDIT A/B MATRIX  (rebuilder vs transformer)")
    print("═" * width)
    print(col)
    print("  " + "─" * (width - 2))

    for name, (build_fn, ai) in SCENARIOS.items():
        # Rebuilder
        orig_r, out_r, audit_r = _run_pipeline(build_fn, ai_results=ai)
        _cleanup(orig_r, out_r)

        # Transformer (use same ai_results if any)
        ai_t = ai if ai else {}
        orig_t = Path(tempfile.mktemp(suffix=".docx"))
        build_fn(orig_t)
        out_t = Path(tempfile.mktemp(suffix=".docx"))
        if out_t.exists(): out_t.unlink()

        from app.services.docx.transformer import transform_docx as _tx
        structure_t     = parse_docx(orig_t)
        style_profile_t = extract_style_profile(orig_t)
        _tx(original_path=orig_t, output_path=out_t, structure=structure_t,
            ai_results=ai_t, style_profile=style_profile_t)
        orig_snap_t = snapshot_document(orig_t, label="original")
        out_snap_t  = snapshot_document(out_t,  label="output")
        audit_t = diff_snapshots(orig_snap_t, out_snap_t, ai_sections=set(ai_t.keys()))
        _cleanup(orig_t, out_t)

        def _row(engine: str, a: "RebuildAudit") -> str:
            return (
                f"  {name:<16} {engine:<13} {SEV.get(a.severity, a.severity):<9} "
                f"{a.para_count_delta:>+6} "
                f"{a.out_image_count - a.orig_image_count:>+5} "
                f"{len(a.headings_missing):>6} "
                f"{a.text_duplications:>4} "
                f"{a.text_truncations:>6} "
                f"{a.style_mismatches:>6}"
            )

        print(_row("rebuilder", audit_r))
        print(_row("transformer", audit_t))
        print("  " + "·" * (width - 2))

    print("═" * width)
    print("  Legend: ΔPara=paragraph count delta, ΔImg=image count delta,")
    print("  MissH=missing headings, Dup=text duplications, Trunc=text truncations")
    print("═" * width + "\n")


if __name__ == "__main__":
    cmd = sys.argv[1] if len(sys.argv) > 1 else "matrix"
    if cmd == "matrix":
        run_matrix()
    elif cmd == "report":
        scenario = sys.argv[2] if len(sys.argv) > 2 else "basic"
        run_report(scenario)
    else:
        print(f"Usage: python {sys.argv[0]} [matrix|report [scenario_name]]")
        print(f"Scenarios: {list(SCENARIOS)}")


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — TRANSFORMER TESTS
# ══════════════════════════════════════════════════════════════════════════════

def _run_transformer(build_fn, ai_results: dict | None = None) -> tuple[Path, Path, "RebuildAudit"]:
    """Same as _run_pipeline but uses transform_docx instead of rebuild_docx."""
    from app.services.docx.transformer import transform_docx

    orig_tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    orig_path = Path(orig_tmp.name)
    orig_tmp.close()
    build_fn(orig_path)

    out_tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    out_path = Path(out_tmp.name)
    out_tmp.close()
    out_path.unlink()

    structure     = parse_docx(orig_path)
    style_profile = extract_style_profile(orig_path)

    transform_docx(
        original_path=orig_path,
        output_path=out_path,
        structure=structure,
        ai_results=ai_results or {},
        style_profile=style_profile,
    )

    orig_snap = snapshot_document(orig_path, label="original")
    out_snap  = snapshot_document(out_path,  label="output")
    audit     = diff_snapshots(orig_snap, out_snap, ai_sections=set((ai_results or {}).keys()))

    return orig_path, out_path, audit


class TestTransformerBasic:
    """Transformer: baseline lab report, no AI replacements."""

    def setup_method(self):
        orig, out, self.audit = _run_transformer(build_lab_report, ai_results={})
        self.orig, self.out = orig, out

    def teardown_method(self):
        _cleanup(self.orig, self.out)

    def test_headings_preserved(self):
        assert self.audit.headings_missing == []

    def test_no_para_delta(self):
        assert self.audit.para_count_delta == 0

    def test_no_duplication(self):
        assert self.audit.text_duplications == 0

    def test_no_truncation(self):
        assert self.audit.text_truncations == 0

    def test_severity_ok(self):
        assert self.audit.severity == "ok"


class TestTransformerAiInjection:
    """Transformer: AI content replaces specific sections."""

    def setup_method(self):
        ai = {
            "1. Aim":    "AI-generated aim: deploy containerised microservices using Kubernetes.",
            "4. Result": "AI-generated result: all pods reached Running state within 90 seconds.",
        }
        orig, out, self.audit = _run_transformer(build_lab_report, ai_results=ai)
        self.orig, self.out, self.ai = orig, out, ai

    def teardown_method(self):
        _cleanup(self.orig, self.out)

    def test_headings_still_present(self):
        assert self.audit.headings_missing == []

    def test_no_duplication(self):
        assert self.audit.text_duplications == 0

    def test_preserved_sections_intact(self):
        # Theory, Procedure, Conclusion were NOT replaced
        assert self.audit.text_truncations == 0, (
            f"Preserved sections were truncated: {self.audit.para_count_delta}"
        )

    def test_ai_content_in_output(self):
        rebuilt = DocxDocument(str(self.out))
        all_text = " ".join(p.text for p in rebuilt.paragraphs)
        assert "Kubernetes" in all_text, "AI content not found in output"

    def test_severity_ok(self):
        assert self.audit.severity in ("ok", "minor"), (
            f"Unexpected severity: {self.audit.severity}"
        )


class TestTransformerImageScenario:
    """Transformer: images stay in-place (the case that broke the rebuilder)."""

    def setup_method(self):
        orig, out, self.audit = _run_transformer(build_with_image, ai_results={
            "Result": "AI result: image displayed correctly in the output document."
        })
        self.orig, self.out = orig, out

    def teardown_method(self):
        _cleanup(self.orig, self.out)

    def test_image_preserved(self):
        assert self.audit.images_lost == 0, (
            f"Image lost: orig={self.audit.orig_image_count} out={self.audit.out_image_count}"
        )

    def test_image_count_unchanged(self):
        assert self.audit.orig_image_count == self.audit.out_image_count

    def test_no_duplication(self):
        assert self.audit.text_duplications == 0, (
            f"Duplication: {self.audit.text_duplications} paragraphs affected"
        )

    def test_no_style_mismatches(self):
        assert self.audit.style_mismatches == 0, (
            f"Style mismatches: {self.audit.style_mismatches}"
        )

    def test_severity_ok(self):
        assert self.audit.severity == "ok", (
            f"Severity: {self.audit.severity}\n"
            + json.dumps(audit_to_dict(self.audit), indent=2)
        )


class TestTransformerTableScenario:
    """Transformer: tables survive and stay in position."""

    def setup_method(self):
        orig, out, self.audit = _run_transformer(build_with_table, ai_results={
            "Conclusion": "AI conclusion: the data in the table confirms the hypothesis."
        })
        self.orig, self.out = orig, out

    def teardown_method(self):
        _cleanup(self.orig, self.out)

    def test_table_not_lost(self):
        assert self.audit.tables_lost == 0

    def test_table_count_matches(self):
        assert self.audit.orig_table_count == self.audit.out_table_count

    def test_no_duplication(self):
        assert self.audit.text_duplications == 0


class TestTransformerVsRebuilder:
    """
    Direct A/B comparison: transformer must match or beat rebuilder
    on the image scenario (the known rebuilder failure case).
    """

    def test_image_scenario_transformer_wins(self):
        """Transformer severity <= rebuilder severity for image documents."""
        _, out_r, audit_r = _run_pipeline(build_with_image, ai_results={})
        _, out_t, audit_t = _run_transformer(build_with_image, ai_results={
            "Result": "AI result content."
        })

        sev_rank = {"ok": 0, "minor": 1, "major": 2, "critical": 3}
        rebuilder_rank   = sev_rank.get(audit_r.severity, 99)
        transformer_rank = sev_rank.get(audit_t.severity, 99)

        assert transformer_rank <= rebuilder_rank, (
            f"Transformer ({audit_t.severity}) worse than rebuilder ({audit_r.severity})"
        )

        # Specifically: transformer should not have image positional drift
        assert audit_t.text_duplications == 0, (
            f"Transformer has duplication (rebuilder bug was duplication)"
        )

        _cleanup(out_r, out_t)
