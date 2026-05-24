"""
tests/test_docx_pipeline.py

Integration tests for the DOCX processing pipeline.
Creates a real DOCX in-memory, writes to temp file, then runs
parser + style_extractor + rebuilder against it.

Run with:
    pytest tests/test_docx_pipeline.py -v
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest
import tempfile

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.docx.parser import parse_docx
from app.services.docx.style_extractor import extract_style_profile
from app.services.docx.rebuilder import rebuild_docx, validate_output


# ── Fixtures ───────────────────────────────────────────────────────────────────

def create_sample_docx(path: Path) -> None:
    """Create a minimal VTU-style lab report DOCX for testing."""
    doc = DocxDocument()

    # Title
    doc.add_heading("BCSL657D — DevOps Lab Report", level=0)

    # Section 1
    doc.add_heading("1. Aim", level=1)
    doc.add_paragraph(
        "To study and implement a CI/CD pipeline using Jenkins and Docker."
    )

    # Section 2
    doc.add_heading("2. Theory", level=1)
    doc.add_paragraph(
        "Continuous Integration (CI) is a development practice where developers "
        "integrate code into a shared repository frequently."
    )
    doc.add_paragraph(
        "Continuous Deployment (CD) extends CI by automatically deploying all "
        "code changes to a testing or production environment after the build stage."
    )

    # Section 3
    doc.add_heading("3. Procedure", level=1)
    doc.add_heading("3.1 Setup", level=2)
    doc.add_paragraph("Install Jenkins on Ubuntu using apt-get.")
    doc.add_heading("3.2 Pipeline Configuration", level=2)
    doc.add_paragraph("Create a Jenkinsfile in the project root.")

    # Section 4
    doc.add_heading("4. Result", level=1)
    doc.add_paragraph(
        "The CI/CD pipeline was successfully configured and tested."
    )

    # Section 5
    doc.add_heading("5. Conclusion", level=1)
    doc.add_paragraph(
        "The experiment demonstrated the importance of automated pipelines."
    )

    doc.save(str(path))


# ── Tests ──────────────────────────────────────────────────────────────────────

class TestParser:
    def setup_method(self):
        self.tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        self.path = Path(self.tmp.name)
        create_sample_docx(self.path)

    def teardown_method(self):
        self.path.unlink(missing_ok=True)

    def test_parse_returns_dict(self):
        result = parse_docx(self.path)
        assert isinstance(result, dict)

    def test_sections_detected(self):
        result = parse_docx(self.path)
        sections = result["sections"]
        assert len(sections) >= 5, f"Expected 5+ sections, got {len(sections)}"

    def test_heading_levels(self):
        result = parse_docx(self.path)
        levels = {s["heading_level"] for s in result["sections"]}
        assert 1 in levels, "Should detect h1 sections"
        assert 2 in levels, "Should detect h2 sub-sections"

    def test_heading_text_preserved(self):
        result = parse_docx(self.path)
        headings = {s["heading_text"] for s in result["sections"]}
        assert "1. Aim" in headings
        assert "4. Result" in headings

    def test_paragraphs_under_sections(self):
        result = parse_docx(self.path)
        aim_section = next(
            s for s in result["sections"] if s["heading_text"] == "1. Aim"
        )
        assert len(aim_section["paragraphs"]) >= 1

    def test_total_paragraphs_positive(self):
        result = parse_docx(self.path)
        assert result["total_paragraphs"] > 0

    def test_title_extracted(self):
        result = parse_docx(self.path)
        # Title heading (level 0) should populate title field
        assert result["title"] != "Untitled"


class TestStyleExtractor:
    def setup_method(self):
        self.tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        self.path = Path(self.tmp.name)
        create_sample_docx(self.path)

    def teardown_method(self):
        self.path.unlink(missing_ok=True)

    def test_returns_dict(self):
        profile = extract_style_profile(self.path)
        assert isinstance(profile, dict)

    def test_page_margins_present(self):
        profile = extract_style_profile(self.path)
        margins = profile.get("page_margins", {})
        assert "top" in margins
        assert "left" in margins
        assert margins["top"] > 0

    def test_heading_styles_extracted(self):
        profile = extract_style_profile(self.path)
        heading_styles = profile.get("heading_styles", {})
        assert "h1" in heading_styles or "h2" in heading_styles

    def test_page_size_present(self):
        profile = extract_style_profile(self.path)
        page_size = profile.get("page_size", {})
        assert "width_pt" in page_size
        assert page_size["width_pt"] > 400  # A4 width ~595pt


class TestRebuilder:
    def setup_method(self):
        self.tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        self.original = Path(self.tmp.name)
        create_sample_docx(self.original)

        self.out_tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        self.output = Path(self.out_tmp.name)
        self.output.unlink()  # rebuilder will create it

    def teardown_method(self):
        self.original.unlink(missing_ok=True)
        self.output.unlink(missing_ok=True)

    def test_rebuild_with_ai_results(self):
        structure = parse_docx(self.original)
        style_profile = extract_style_profile(self.original)

        ai_results = {
            "1. Aim": "To implement and study a Maven build lifecycle using POM configuration.",
            "4. Result": "The Maven build was successful. All lifecycle phases completed without errors.",
        }

        result_path = rebuild_docx(
            original_path=self.original,
            output_path=self.output,
            structure=structure,
            ai_results=ai_results,
            style_profile=style_profile,
        )

        assert result_path.exists()
        assert result_path.stat().st_size > 1000  # not empty

    def test_output_validates(self):
        structure = parse_docx(self.original)
        style_profile = extract_style_profile(self.original)

        rebuild_docx(
            original_path=self.original,
            output_path=self.output,
            structure=structure,
            ai_results={"1. Aim": "Updated aim text."},
            style_profile=style_profile,
        )

        assert validate_output(self.output) is True

    def test_rebuild_preserves_untouched_sections(self):
        structure = parse_docx(self.original)
        style_profile = extract_style_profile(self.original)

        # Only regenerate Aim — Conclusion should be preserved from original
        rebuild_docx(
            original_path=self.original,
            output_path=self.output,
            structure=structure,
            ai_results={"1. Aim": "New aim content."},
            style_profile=style_profile,
        )

        rebuilt = DocxDocument(str(self.output))
        all_text = " ".join(p.text for p in rebuilt.paragraphs)

        # Conclusion's original text should be preserved
        assert "automated pipelines" in all_text

    def test_rebuild_replaces_ai_sections(self):
        structure = parse_docx(self.original)
        style_profile = extract_style_profile(self.original)

        rebuild_docx(
            original_path=self.original,
            output_path=self.output,
            structure=structure,
            ai_results={"1. Aim": "REPLACED_AIM_CONTENT_XYZ"},
            style_profile=style_profile,
        )

        rebuilt = DocxDocument(str(self.output))
        all_text = " ".join(p.text for p in rebuilt.paragraphs)
        assert "REPLACED_AIM_CONTENT_XYZ" in all_text


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
