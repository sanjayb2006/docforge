"""
tests/test_transformer.py

Targeted tests for the transformer fallback behavior, especially when
sections carry a paragraph-level section break.
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document as DocxDocument
from docx.enum.section import WD_SECTION

from app.services.docx.parser import parse_docx
from app.services.docx.transformer import transform_docx


def create_section_break_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("1. Aim", level=1)
    doc.add_paragraph("Paragraph 1 of section 1.")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("2. Theory", level=1)
    doc.add_paragraph("Paragraph 1 of section 2.")
    doc.save(str(path))


def test_transform_preserves_section_break_paragraph() -> None:
    tmp_orig = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    original = Path(tmp_orig.name)
    tmp_orig.close()

    output_path = original.parent / (original.stem + "_out.docx")

    try:
        create_section_break_docx(original)
        structure = parse_docx(original)

        transform_docx(
            original_path=original,
            output_path=output_path,
            structure=structure,
            ai_results={
                "1. Aim": (
                    "This is rewritten first paragraph.\n\n"
                    "This is newly inserted overflow paragraph"
                ),
            },
        )

        rebuilt = DocxDocument(str(output_path))
        texts = [p.text for p in rebuilt.paragraphs]

        assert "This is rewritten first paragraph." in texts
        assert "This is newly inserted overflow paragraph" in texts

        first_heading_index = texts.index("2. Theory")
        overflow_index = texts.index("This is newly inserted overflow paragraph")
        assert overflow_index < first_heading_index
    finally:
        original.unlink(missing_ok=True)
        output_path.unlink(missing_ok=True)
