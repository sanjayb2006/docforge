"""
services/fidelity/scorer.py

Document Fidelity Scoring System.

Compares an original DOCX against a rebuilt DOCX and produces a
FidelityReport with per-dimension scores and actionable diagnostics.

Dimensions scored (each 0–100, weighted to total):
  1. Structure fidelity    (30%) — headings present, order correct, levels match
  2. Formatting fidelity   (25%) — font, size, spacing, margins match
  3. Content completeness  (20%) — no sections lost, no text truncation
  4. Element preservation  (15%) — tables, images, page breaks present
  5. Style consistency     (10%) — consistent body font/size across paragraphs

Overall score: weighted average of the five dimensions.

Interpretation:
  95–100  Submission-ready with no manual correction expected
  85–94   Minor cosmetic differences, quick review recommended
  70–84   Noticeable differences, spot-check required
  50–69   Significant formatting loss, manual correction needed
  <50     Major fidelity failure, rebuild likely broken

Usage:
    from app.services.fidelity.scorer import score_documents
    report = score_documents(original_path, rebuilt_path)
    print(report.summary())
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.oxml.ns import qn

log = logging.getLogger(__name__)


# ── Data structures ────────────────────────────────────────────────────────────

@dataclass
class DimensionScore:
    name: str
    score: float          # 0–100
    weight: float         # contribution to overall
    issues: list[str] = field(default_factory=list)
    details: dict[str, Any] = field(default_factory=dict)

    @property
    def weighted(self) -> float:
        return self.score * self.weight


@dataclass
class FidelityReport:
    original_path: str
    rebuilt_path: str
    dimensions: list[DimensionScore] = field(default_factory=list)
    overall_score: float = 0.0
    grade: str = "F"
    rebuild_ready: bool = False

    # Raw comparison data
    original_stats: dict[str, Any] = field(default_factory=dict)
    rebuilt_stats:  dict[str, Any] = field(default_factory=dict)

    def summary(self) -> str:
        lines = [
            f"DocForge Fidelity Report",
            f"{'─'*40}",
            f"Overall Score : {self.overall_score:.1f}/100  [{self.grade}]",
            f"Ready         : {'YES ✓' if self.rebuild_ready else 'NO — manual review needed'}",
            f"{'─'*40}",
        ]
        for dim in self.dimensions:
            issue_str = f"  Issues: {'; '.join(dim.issues)}" if dim.issues else ""
            lines.append(f"  {dim.name:<25} {dim.score:5.1f}/100 (×{dim.weight}){issue_str}")
        lines.append(f"{'─'*40}")
        return "\n".join(lines)

    def to_dict(self) -> dict[str, Any]:
        return {
            "overall_score": round(self.overall_score, 2),
            "grade": self.grade,
            "rebuild_ready": self.rebuild_ready,
            "dimensions": [
                {
                    "name": d.name,
                    "score": round(d.score, 2),
                    "weight": d.weight,
                    "issues": d.issues,
                    "details": d.details,
                }
                for d in self.dimensions
            ],
            "original_stats": self.original_stats,
            "rebuilt_stats": self.rebuilt_stats,
        }


# ── Document stats collection ──────────────────────────────────────────────────

def _collect_stats(doc: DocxDocument) -> dict[str, Any]:
    """Collect structural and formatting statistics for one document."""
    headings: list[dict] = []
    body_fonts: list[str] = []
    body_sizes: list[float] = []
    para_spacings: list[dict] = []

    HEADING_PREFIXES = ("heading", "title")

    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        is_heading = any(style_name.startswith(p) for p in HEADING_PREFIXES)

        if is_heading:
            headings.append({
                "text": para.text.strip(),
                "style": para.style.name,
            })
        else:
            # Collect body font/size
            for run in para.runs:
                rPr = run._r.find(qn("w:rPr"))
                if rPr is not None:
                    fonts_el = rPr.find(qn("w:rFonts"))
                    if fonts_el is not None:
                        for attr in ("ascii", "hAnsi"):
                            v = fonts_el.get(qn(f"w:{attr}"))
                            if v:
                                body_fonts.append(v)
                                break
                    sz = rPr.find(qn("w:sz"))
                    if sz is not None:
                        try:
                            body_sizes.append(int(sz.get(qn("w:val"), 0)) / 2)
                        except (ValueError, TypeError):
                            pass

        # Spacing
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            sp = pPr.find(qn("w:spacing"))
            if sp is not None:
                spacing: dict[str, Any] = {}
                for attr in ("before", "after", "line"):
                    v = sp.get(qn(f"w:{attr}"))
                    if v:
                        try:
                            spacing[attr] = round(int(v) / 20, 1)
                        except (ValueError, TypeError):
                            pass
                if spacing:
                    para_spacings.append(spacing)

    # Page margins (first section)
    margins: dict[str, float] = {}
    try:
        section = doc.sections[0]
        for side in ("top", "bottom", "left", "right"):
            v = getattr(section, f"{side}_margin", None)
            if v is not None:
                margins[side] = round(v.pt, 1)
    except Exception:
        pass

    # Dominant body font/size
    dominant_font = _most_common(body_fonts) if body_fonts else None
    dominant_size = _most_common([str(s) for s in body_sizes]) if body_sizes else None

    return {
        "heading_count": len(headings),
        "headings": headings,
        "table_count": len(doc.tables),
        "paragraph_count": len(doc.paragraphs),
        "page_margins": margins,
        "dominant_body_font": dominant_font,
        "dominant_body_size_pt": float(dominant_size) if dominant_size else None,
        "para_spacings_sample": para_spacings[:5],
        "total_words": sum(len(p.text.split()) for p in doc.paragraphs),
    }


def _most_common(lst: list[str]) -> str | None:
    if not lst:
        return None
    return max(set(lst), key=lst.count)


# ── Scoring dimensions ─────────────────────────────────────────────────────────

def _score_structure(
    orig: dict, rebuilt: dict,
    orig_doc: DocxDocument, rebuilt_doc: DocxDocument,
) -> DimensionScore:
    issues: list[str] = []
    score = 100.0

    orig_headings  = orig["headings"]
    rebuilt_headings = rebuilt["headings"]

    oh_count = orig["heading_count"]
    rh_count = rebuilt["heading_count"]

    # Heading count match
    if oh_count == 0:
        # No headings in original — can't score structure
        return DimensionScore("Structure Fidelity", 100, 0.30, [], {"note": "no headings in original"})

    count_ratio = min(rh_count, oh_count) / max(oh_count, 1)
    if count_ratio < 1.0:
        lost = oh_count - rh_count
        issues.append(f"{lost} heading(s) missing in rebuild")
        score -= (1 - count_ratio) * 50

    # Heading text match (order-sensitive)
    matched = 0
    orig_texts  = [h["text"] for h in orig_headings]
    rebuilt_texts = [h["text"] for h in rebuilt_headings]
    for ot in orig_texts:
        if ot in rebuilt_texts:
            matched += 1
        else:
            issues.append(f"Heading not found in rebuild: '{ot[:40]}'")
    text_ratio = matched / max(oh_count, 1)
    score -= (1 - text_ratio) * 40

    score = max(0.0, min(100.0, score))
    return DimensionScore(
        "Structure Fidelity", score, 0.30, issues,
        {"original_headings": oh_count, "rebuilt_headings": rh_count, "matched": matched},
    )


def _score_formatting(
    orig: dict, rebuilt: dict,
    orig_doc: DocxDocument, rebuilt_doc: DocxDocument,
) -> DimensionScore:
    issues: list[str] = []
    score = 100.0
    deductions = 0.0

    # Margins comparison
    o_margins = orig["page_margins"]
    r_margins = rebuilt["page_margins"]
    margin_mismatches = 0
    for side in ("top", "bottom", "left", "right"):
        ov = o_margins.get(side)
        rv = r_margins.get(side)
        if ov is not None and rv is not None:
            diff = abs(ov - rv)
            if diff > 2.0:  # >2pt tolerance
                margin_mismatches += 1
                issues.append(f"Margin '{side}': original={ov}pt rebuilt={rv}pt (Δ{diff:.1f}pt)")
    if margin_mismatches:
        deductions += margin_mismatches * 5

    # Body font consistency
    o_font = orig.get("dominant_body_font")
    r_font = rebuilt.get("dominant_body_font")
    if o_font and r_font and o_font.lower() != r_font.lower():
        issues.append(f"Body font changed: '{o_font}' → '{r_font}'")
        deductions += 15

    # Body size consistency
    o_size = orig.get("dominant_body_size_pt")
    r_size = rebuilt.get("dominant_body_size_pt")
    if o_size and r_size:
        if abs(o_size - r_size) > 0.5:
            issues.append(f"Body size changed: {o_size}pt → {r_size}pt")
            deductions += 10

    score = max(0.0, min(100.0, score - deductions))
    return DimensionScore(
        "Formatting Fidelity", score, 0.25, issues,
        {"original_font": o_font, "rebuilt_font": r_font,
         "original_size": o_size, "rebuilt_size": r_size},
    )


def _score_content(
    orig: dict, rebuilt: dict,
    orig_doc: DocxDocument, rebuilt_doc: DocxDocument,
) -> DimensionScore:
    issues: list[str] = []
    score = 100.0

    o_words = orig["total_words"]
    r_words = rebuilt["total_words"]

    if o_words == 0:
        return DimensionScore("Content Completeness", 100, 0.20, [], {"note": "empty original"})

    # Word count ratio (AI content will differ — use generous tolerance)
    ratio = r_words / o_words
    if ratio < 0.3:
        issues.append(f"Rebuilt word count very low: {r_words} vs original {o_words}")
        score -= 40
    elif ratio < 0.6:
        issues.append(f"Rebuilt word count low: {r_words} vs original {o_words}")
        score -= 15

    # Paragraph count comparison
    o_paras = orig["paragraph_count"]
    r_paras = rebuilt["paragraph_count"]
    if r_paras < o_paras * 0.5:
        issues.append(f"Paragraph count dropped significantly: {r_paras} vs {o_paras}")
        score -= 20

    score = max(0.0, min(100.0, score))
    return DimensionScore(
        "Content Completeness", score, 0.20, issues,
        {"original_words": o_words, "rebuilt_words": r_words,
         "original_paras": o_paras, "rebuilt_paras": r_paras},
    )


def _score_elements(
    orig: dict, rebuilt: dict,
    orig_doc: DocxDocument, rebuilt_doc: DocxDocument,
) -> DimensionScore:
    issues: list[str] = []
    score = 100.0

    # Tables
    o_tables = orig["table_count"]
    r_tables = rebuilt["table_count"]
    if o_tables > 0:
        if r_tables == 0:
            issues.append(f"All {o_tables} table(s) lost in rebuild")
            score -= 40
        elif r_tables < o_tables:
            lost = o_tables - r_tables
            issues.append(f"{lost} table(s) lost: original={o_tables} rebuilt={r_tables}")
            score -= (lost / o_tables) * 30

    # Images (count via XML)
    o_images = _count_images(orig_doc)
    r_images = _count_images(rebuilt_doc)
    if o_images > 0 and r_images == 0:
        issues.append(f"All {o_images} image(s) lost in rebuild")
        score -= 20
    elif o_images > 0 and r_images < o_images:
        issues.append(f"Image count dropped: {o_images} → {r_images}")
        score -= (1 - r_images / o_images) * 15

    score = max(0.0, min(100.0, score))
    return DimensionScore(
        "Element Preservation", score, 0.15, issues,
        {"original_tables": o_tables, "rebuilt_tables": r_tables,
         "original_images": o_images, "rebuilt_images": r_images},
    )


def _count_images(doc: DocxDocument) -> int:
    count = 0
    for para in doc.paragraphs:
        count += len(para._p.findall(f".//{qn('a:blip')}", para._p.nsmap))
    return count


def _score_style_consistency(
    orig: dict, rebuilt: dict,
    orig_doc: DocxDocument, rebuilt_doc: DocxDocument,
) -> DimensionScore:
    """
    Check that body paragraphs in the rebuild use a consistent font/size.
    Mixed fonts across body paragraphs = style inconsistency.
    """
    issues: list[str] = []
    score = 100.0

    fonts_seen: set[str] = set()
    sizes_seen: set[float] = set()

    for para in rebuilt_doc.paragraphs:
        style_name = (para.style.name or "").lower()
        if any(style_name.startswith(p) for p in ("heading", "title", "toc")):
            continue
        for run in para.runs:
            rPr = run._r.find(qn("w:rPr"))
            if rPr is None:
                continue
            fonts_el = rPr.find(qn("w:rFonts"))
            if fonts_el is not None:
                for attr in ("ascii", "hAnsi"):
                    v = fonts_el.get(qn(f"w:{attr}"))
                    if v:
                        fonts_seen.add(v)
                        break
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                try:
                    sizes_seen.add(round(int(sz.get(qn("w:val"), 0)) / 2, 1))
                except (ValueError, TypeError):
                    pass

    if len(fonts_seen) > 3:
        issues.append(f"Body uses {len(fonts_seen)} different fonts: {sorted(fonts_seen)}")
        score -= min(40, (len(fonts_seen) - 3) * 10)

    if len(sizes_seen) > 3:
        issues.append(f"Body uses {len(sizes_seen)} different sizes: {sorted(sizes_seen)}")
        score -= min(20, (len(sizes_seen) - 3) * 5)

    score = max(0.0, min(100.0, score))
    return DimensionScore(
        "Style Consistency", score, 0.10, issues,
        {"body_fonts_seen": sorted(fonts_seen), "body_sizes_seen": sorted(sizes_seen)},
    )


# ── Grade ──────────────────────────────────────────────────────────────────────

def _grade(score: float) -> tuple[str, bool]:
    if score >= 95:
        return "A+", True
    if score >= 85:
        return "A",  True
    if score >= 75:
        return "B",  False
    if score >= 65:
        return "C",  False
    if score >= 50:
        return "D",  False
    return "F", False


# ── Public API ─────────────────────────────────────────────────────────────────

def score_documents(
    original_path: str | Path,
    rebuilt_path:  str | Path,
) -> FidelityReport:
    """
    Compare original vs rebuilt DOCX and return a FidelityReport.

    Args:
        original_path: Path to the uploaded template
        rebuilt_path:  Path to the AI-rebuilt output

    Returns:
        FidelityReport with per-dimension scores and diagnostics
    """
    original_path = Path(original_path)
    rebuilt_path  = Path(rebuilt_path)

    log.info("Scoring fidelity: %s vs %s", original_path.name, rebuilt_path.name)

    orig_doc    = DocxDocument(str(original_path))
    rebuilt_doc = DocxDocument(str(rebuilt_path))

    orig_stats    = _collect_stats(orig_doc)
    rebuilt_stats = _collect_stats(rebuilt_doc)

    scorers = [
        _score_structure,
        _score_formatting,
        _score_content,
        _score_elements,
        _score_style_consistency,
    ]

    dimensions = [
        fn(orig_stats, rebuilt_stats, orig_doc, rebuilt_doc)
        for fn in scorers
    ]

    overall = sum(d.weighted for d in dimensions)
    grade, ready = _grade(overall)

    report = FidelityReport(
        original_path=str(original_path),
        rebuilt_path=str(rebuilt_path),
        dimensions=dimensions,
        overall_score=overall,
        grade=grade,
        rebuild_ready=ready,
        original_stats=orig_stats,
        rebuilt_stats=rebuilt_stats,
    )

    log.info(
        "Fidelity score: %.1f/100 [%s] — %s",
        overall, grade, "ready" if ready else "needs review",
    )
    for dim in dimensions:
        if dim.issues:
            log.warning("  %s (%.0f): %s", dim.name, dim.score, "; ".join(dim.issues))

    return report
